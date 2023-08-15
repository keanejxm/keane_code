#!/usr/bin/python3
# -*- coding:utf-8 -*-
"""
:author: keane
:file  operate_docx.py
:time  2023/7/6 10:07
:desc  
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import pandas as pd


class Header:
    def __init__(self, value, column_index):
        self.value = value
        self.column_index = column_index


class RowIndex:
    def __init__(self, value, row_index):
        self.value = value
        self.row_index = row_index


class DealDocxTable:
    def __init__(self, table, header=0, index=0):
        self._table = table
        self._header = header
        self._index = index
        self.headers = self._get_headers(header)
        self.rows_index = self._get_rows_index(index)
        self.headers_values = [header.value for header in self.headers]
        self.rows_index_values = [row.value for row in self.rows_index]

    def _get_headers(self, row_index=0) -> list:
        """获取表头"""
        headers = list()
        for column_index, column in enumerate(self._table.columns):
            if row_index is None:
                headers.append(Header(column_index, column_index))
            else:
                cell_text = self._table.cell(row_index, column_index).text
                headers.append(Header(cell_text, column_index))
        return headers

    def _get_rows_index(self, column_index=0):
        """获取列名"""
        rows_index = list()
        for row_index, row in enumerate(self._table.rows):
            if column_index is None:
                rows_index.append(RowIndex(row_index, row_index))
            else:
                cell_text = self._table.cell(row_index, column_index).text
                rows_index.append(RowIndex(cell_text, row_index))
        return rows_index

    def update_data(self, row, column, update_value):
        """根据列名插入数据"""
        row_index, column_index = self._find_index(row, column)
        self._table.cell(row_index, column_index).text = update_value
        for para in self._table.cell(row_index, column_index).paragraphs:
            for run in para.runs:
                run.font.name = '楷体'
                run.text = update_value
                run.font.size = Pt(12)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    def insert_row(self):
        """插入一行"""
        self._table.add_row()
        # 插入一行后重新重置行和列
        self.headers = self._get_headers(self._header)
        self.rows_index = self._get_rows_index(self._index)
        self.headers_values = [header.value for header in self.headers]
        self.rows_index_values = [row.value for row in self.rows_index]

    def _find_index(self, row, column):
        """查询索引"""
        if isinstance(column, str):
            column_index = [header.column_index for header in self.headers if header.value == column]
            assert column_index, f"无【{column}】索引"
            column_index = column_index[0]
        elif isinstance(column, int):
            column_index = column
        else:
            raise IndexError("错误的索引")
        if isinstance(row, str):
            row_index = [row1.row_index for row1 in self.rows_index if row1.value == row]
            assert row_index, f"无【{row}】索引"
            row_index = row_index[0]
        elif isinstance(row, int):
            row_index = row
        else:
            raise IndexError("错误的索引")
        return row_index, column_index

    def search_index(self, value):
        """根据值查询行"""
        for row_index in self.rows_index_values:
            for header_index in self.headers_values:
                cell_text = self.get_value(row_index, header_index)
                if cell_text == value:
                    return row_index, header_index
        else:
            return None, None

    def get_value(self, row, column):
        """获取docx中表格数据"""
        row_index, column_index = self._find_index(row, column)
        cell_text = self._table.cell(row_index, column_index).text
        return cell_text

    def table_to_df(self):
        """将表格转为dataFrame,此方法更容易提取table中的数据"""
        table_list = list()
        column_max = 0
        for row in self._table.rows:
            row_list = list()
            for cell in row.cells:
                if cell.text not in row_list:
                    row_list.append(cell.text)
            column_max = len(row_list) if len(row_list) > column_max else column_max
            if row_list not in table_list:
                table_list.append(row_list)
        if self._header is None:
            column = [f"column{i}" for i in range(column_max)]
        else:
            column = table_list.pop(self._header)
            for i in range(column_max - len(column)):
                column.append(f"column_{len(column) + i}")
        df = pd.DataFrame(table_list, columns=column)
        return df

    def remove_table_tr(self):
        doc = Document("建信科创新兴一号私募股权投资基金.docx")
        for table in doc.tables:
            table_cell_text = table.cell(0, 0).text
            if table_cell_text == "项目":
                deal_doc_table = DealDocxTable(table, index=1)
                if "报告期期间" in deal_doc_table.headers_values:
                    tbl = table._tbl
                    for row in table.rows:
                        row_value = list()
                        for cell in row.cells:
                            row_value.append(cell.text)
                        print(row_value)
                        tbl.remove(row._tr)
                        break
        doc.save("aaa.docx")
    def number_deal(self):
        num_value = 1000000
        if "." in str(num_value):
            # nums:整数位，surplus:小数位
            nums, surplus = str(num_value).split(".")
        else:
            nums, surplus = str(num_value), "0"
        res = ""
        num = 0
        for word in reversed(nums):
            res += word
            num += 1
            if num == 3:
                res += ","
                num = 0
        if surplus != "0":
            res = "".join(reversed(res)) + "." + surplus
        else:
            res = "".join(reversed(res))
        res = res.lstrip(",")
        print(res)


if __name__ == '__main__':

    doc =Document("8账户业务申请表-产品投资者.docx")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text.split("\n"))
                print("*"*100)
