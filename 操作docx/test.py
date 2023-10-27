#!/usr/bin/python3
# -*- coding:utf-8 -*-
"""
:author: keane
:file  test.py
:time  2023/7/24 17:02
:desc  
"""
import pandas as pd
import numpy as np
from docx import Document
from 操作docx.operate_docx import DealDocxTable

doc1 = Document("开放式基金账户业务申请表(产品).docx")


def table_docx_run(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        print(run.text)


def para_doc_run(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            print(run.text)


class PrintMessage:
    def __init__(self, function_):
        self.function_ = function_

    def __call__(self, *args, **kwargs):
        print(f"装饰器{self.function_()}")


@PrintMessage
def hello_world():
    return "Hello World"


example_a = {
    "aa": ["aaa", "bbbb"],
    "bb": ["ccc", "dddd"]
}

df = pd.DataFrame(example_a)
columns = df.columns
new_df = pd.DataFrame(columns=columns)
filter_df1 = df[df["aa"] == "aaa"]
filter_df2 = df[df["aa"] == "bbbb"]
filter_df1["dd"]="11111"
new_df = pd.concat([new_df,filter_df1])
print(new_df)
new_df =pd.concat((new_df,filter_df2))
print(new_df)
wirter = pd.ExcelWriter("aaa.xlsx",engine="openpyxl")
new_df.to_excel(wirter,index=False)
wirter.close()