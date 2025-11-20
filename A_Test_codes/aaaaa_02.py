#!/usr/bin/python3
# -*- coding:utf-8 -*-
"""
:author: keane
:file  aaaaa_02.py
:time  2025/3/13 9:53
:desc  # 不平衡数据
"""

import re
import jieba
from typing import List

import pandas as pd
from caffe2.python.operator_test.hsm_test import words
from docx import Document

# # 获取文章中的句子
# docx = Document("改进3.21赵彦玲毕业论文 .docx")
#
# article = ""
# for paragraph in docx.paragraphs:
#     article += paragraph.text
#
# import re


def split_sentences(text):
    """
    基于中文句末标点进行简单句子分割
    支持常见中文句末标点：。！？；…
    """
    sentences = re.split(r'(?<=[。！？；…])', text)
    # 去除空字符串和空白字符
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences


# 使用示例
# article = """
# 自然语言处理(NLP)是人工智能的重要领域。它涉及计算机科学、语言学和机器学习等多个学科！
# NLP的应用包括：机器翻译、情感分析、语音识别等。你觉得这些技术有趣吗？
# """
# sentences = split_sentences(article)
# for i, sentence in enumerate(sentences, 1):
#     sentence = sentence.replace(" ","")
#     with open("article.txt", "a", encoding="utf-8") as f:
#         f.write(sentence + "\n")


# 读取文献内容
import os

import fitz  # PyMuPDF


def extract_text_from_pdf(pdf_path):
    """
    从PDF提取文本内容
    """
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text().replace("\n", "")
    return text


# # 使用示例
# pdf_text = extract_text_from_pdf("converted_from_caj.pdf")
# print(pdf_text)
def read_docx(file_path):
    docx = Document(file_path)

    article = ""
    for paragraph in docx.paragraphs:
        article += paragraph.text
    # docx.close()
    return article


wenxian_dict = dict()

for file_name in os.listdir("wenxian"):
    if file_name.endswith(".pdf"):
        content = extract_text_from_pdf(os.path.join("wenxian", file_name))
        wenxian_dict[file_name] = content
    elif file_name.endswith(".docx"):
        content = read_docx(os.path.join("wenxian", file_name))
        wenxian_dict[file_name] = content
    else:
        pass


def judge_sentence(sentence, pdf_content):
    words = jieba.lcut(sentence)

    pdf_sentences = split_sentences(pdf_content)

    for pdf_sentence in pdf_sentences:
        num = 0
        for word in words:
            if word in pdf_sentence:
                num += 1
        if num / len(words) >= 0.7:
            yield sentence, pdf_sentence


# data = dict()
# wx_name_list = list()
# artcle_list = list()
# pdf_list= list()
# with open("article.txt", "r", encoding="utf-8") as f:
#     articles = f.readlines()
#     for article in articles:
#         for wx_name, wenxian_content in wenxian_dict.items():
#             for sentence,pdf_sentence in  judge_sentence(article, wenxian_content):
#                 wx_name_list.append(wx_name)
#                 artcle_list.append(article)
#                 pdf_list.append(pdf_sentence)
#                 print(wx_name,"\n",article,"\n",pdf_sentence)
#                 print("-------------------------------------------------------------------------------------------")
# import pandas as pd
#
# new_data ={"文献名字":wx_name_list,"论文句子":artcle_list,"文献句子":pdf_list}
# df = pd.DataFrame(new_data)
# writer = pd.ExcelWriter("article.xlsx", engine="openpyxl")
# df.to_excel(writer, index=False)
# writer.save()
pdf_path = "改进3.31赵彦玲毕业论文 _20250401224646.pdf"
doc = fitz.open(pdf_path)
text = ""
a = list()
b = list()
for page in doc:
    page_num, page_text = page.number, page.get_text()
    res = re.findall("\[\d+\]",page_text)
    if res:
        a.append(page_num)
        b.append(res)
df = pd.DataFrame({"页码":a,"文献":b})
writer = pd.ExcelWriter("aaa.xlsx", engine="openpyxl")
df.to_excel(writer, sheet_name="Sheet1",index=False)
writer.save()
writer.close()